import os
import json
import subprocess
import shutil
import qrcode
import tempfile
import boto3
import re
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from api.extensions import db
from api.models.im_certificates import IMCertificate
from api.models.instructionalmaterials import InstructionalMaterial
from api.models.authors import Author
from api.models.users import User
from api.services.email_service import EmailService

class CertificateService:
    TEMPLATE_S3_KEY = 'requirements/cert-of-appreciation.docx'
    GENERATED_CERTIFICATES_PREFIX = 'generated-certificates'
    QR_PLACEHOLDER = '[QR CODE SPACE]'
    QR_INLINE_WIDTH_INCHES = 1.5
    QR_FLOATING_WIDTH_INCHES = 1.35

    @staticmethod
    def generate_certificates(im_id, template_path=None):
        """Generate certificates for all authors of an IM.
        
        Args:
            im_id: The instructional material ID.
            template_path: Optional path to a custom DOCX template. If not
                provided the default template is downloaded from S3.
        """
        im = InstructionalMaterial.query.get(im_id)
        if not im:
            raise ValueError("Instructional Material not found")
        
        # Get all authors
        authors = Author.query.filter_by(im_id=im_id).all()
        if not authors:
            raise ValueError("No authors found for this IM")
        
        # Get IM details
        college_name, course_code, course_title, program_name = CertificateService._get_im_details(im)
        semester = im.semester or "N/A"
        academic_year = CertificateService._format_academic_year(im.validity)
        date_issued = date.today().strftime("%B %d, %Y")
        
        # Download template from S3 only if a custom one wasn't supplied
        owns_template = template_path is None
        if owns_template:
            template_path = CertificateService._download_template()
        
        certificates = []
        try:
            for author in authors:
                user = User.query.get(author.user_id)
                if not user:
                    continue
                
                # Generate certificate
                cert_data = CertificateService._generate_certificate(
                    template_path, user, college_name, course_code, course_title,
                    program_name, semester, academic_year, date_issued, im_id, im.validity
                )
                
                certificates.append(cert_data)
        finally:
            # Cleanup template only if we downloaded it
            if owns_template and os.path.exists(template_path):
                os.remove(template_path)
        
        return certificates
    
    @staticmethod
    def generate_certificate_for_user(im_id, user_id, template_path=None):
        """Generate and send a certificate for a single author only.
        
        Useful for post-publish catch-up when an author was missed.
        """
        im = InstructionalMaterial.query.get(im_id)
        if not im:
            raise ValueError("Instructional Material not found")
        
        user = User.query.get(user_id)
        if not user:
            raise ValueError("User not found")
        
        college_name, course_code, course_title, program_name = CertificateService._get_im_details(im)
        semester = im.semester or "N/A"
        academic_year = CertificateService._format_academic_year(im.validity)
        date_issued = date.today().strftime("%B %d, %Y")
        
        owns_template = template_path is None
        if owns_template:
            template_path = CertificateService._download_template()
        
        try:
            cert_data = CertificateService._generate_certificate(
                template_path, user, college_name, course_code, course_title,
                program_name, semester, academic_year, date_issued, im_id, im.validity
            )
        finally:
            if owns_template and os.path.exists(template_path):
                os.remove(template_path)
        
        return cert_data

    @staticmethod
    def get_certificates_for_user(user_id):
        """Return all certificates issued to a user, enriched with IM details."""
        from api.models.im_certificates import IMCertificate
        from api.models.instructionalmaterials import InstructionalMaterial
        
        certs = IMCertificate.query.filter_by(user_id=user_id).order_by(IMCertificate.date_issued.desc()).all()
        result = []
        for cert in certs:
            im = InstructionalMaterial.query.get(cert.im_id)
            college_name, course_code, course_title, _ = CertificateService._get_im_details(im) if im else ('N/A', 'N/A', 'N/A', 'N/A')
            result.append({
                'id': cert.id,
                'qr_id': cert.qr_id,
                'im_id': cert.im_id,
                'user_id': cert.user_id,
                's3_link': CertificateService._try_presign_pdf(cert.qr_id),
                's3_link_docx': CertificateService._resolve_docx_link(cert.qr_id),
                'date_issued': cert.date_issued.isoformat() if cert.date_issued else None,
                'created_at': cert.created_at.isoformat() if cert.created_at else None,
                'subject_code': course_code,
                'subject_title': course_title,
                'im_version': im.version if im else None,
            })
        return result

    @staticmethod
    def _get_im_details(im):
        """Extract college, course, and program details from IM"""
        if im.university_im_id and im.university_im:
            uni = im.university_im
            college_name = uni.college.name if uni.college else "N/A"
            course_code = uni.subject.code if uni.subject else "N/A"
            course_title = uni.subject.name if uni.subject else "N/A"
            program_name = uni.department.name if uni.department else "N/A"
        elif im.service_im_id and im.service_im:
            svc = im.service_im
            college_name = svc.college.name if svc.college else "N/A"
            course_code = svc.subject.code if svc.subject else "N/A"
            course_title = svc.subject.name if svc.subject else "N/A"
            program_name = "Service Course"
        else:
            college_name = course_code = course_title = program_name = "N/A"
        
        return college_name, course_code, course_title, program_name
    
    @staticmethod
    def _format_academic_year(validity):
        """Convert validity to academic year format (e.g., 2025 -> 2025-2026)"""
        try:
            year = int(validity)
            return f"{year}-{year + 1}"
        except:
            return validity
    
    @staticmethod
    def _download_template():
        """Download certificate template from S3"""
        bucket_name = os.getenv('AWS_BUCKET_NAME')
        s3_key = CertificateService.TEMPLATE_S3_KEY
        
        s3 = boto3.client('s3')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.close()
        
        s3.download_file(bucket_name, s3_key, temp_file.name)
        return temp_file.name
    
    @staticmethod
    def _generate_certificate(template_path, user, college_name, course_code, course_title,
                             program_name, semester, academic_year, date_issued, im_id, validity):
        """Generate a single certificate for a user"""
        # Load template
        doc = Document(template_path)

        author_name = CertificateService._build_author_name(user)
        author_rank = user.rank or ""
        author_rank_and_name = f"{author_rank} {author_name}".strip()
        semester_label = CertificateService._format_semester_label(semester)
        validity_duration = CertificateService._format_validity_duration(validity)
        course_code_and_title = f"{course_code}: {course_title}"

        replacements, regex_replacements = CertificateService._build_replacement_maps(
            college_name=college_name,
            course_code=course_code,
            course_title=course_title,
            author_rank=author_rank,
            author_name=author_name,
            author_rank_and_name=author_rank_and_name,
            program_name=program_name,
            semester_label=semester_label,
            academic_year=academic_year,
            date_issued=date_issued,
            course_code_and_title=course_code_and_title,
            validity_duration=validity_duration,
        )

        CertificateService._apply_template_replacements(
            doc,
            replacements=replacements,
            regex_replacements=regex_replacements,
            semester_label=semester_label,
            academic_year=academic_year,
            date_issued=date_issued,
            validity_duration=validity_duration,
        )
        
        # Create certificate record to get ID
        cert = IMCertificate(
            qr_id=f"CERT-TEMP",
            im_id=im_id,
            user_id=user.id,
            s3_link="",
            date_issued=date.today()
        )
        db.session.add(cert)
        db.session.flush()
        
        # Update QR ID with actual ID
        cert.qr_id = f"CERT-{cert.id}"
        
        # Generate QR code
        qr_data = {
            "qr_id": cert.qr_id,
            "author_name": author_name,
            "im_id": im_id,
            "date_issued": date_issued
        }
        qr_img = CertificateService._generate_qr_code(json.dumps(qr_data))
        
        # Add QR code to document (bottom right)
        CertificateService._add_qr_to_document(doc, qr_img)
        
        # Save DOCX to temp file
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_output.name)
        temp_output.close()

        # Convert DOCX → PDF (best-effort; won't crash if unavailable)
        pdf_path = CertificateService._convert_docx_to_pdf(temp_output.name)

        # Upload DOCX to S3
        docx_key = f"{CertificateService.GENERATED_CERTIFICATES_PREFIX}/{cert.qr_id}.docx"
        docx_s3_link = CertificateService._upload_to_s3(temp_output.name, docx_key)

        # Upload PDF to S3 if conversion succeeded
        pdf_s3_link = None
        if pdf_path:
            pdf_key = f"{CertificateService.GENERATED_CERTIFICATES_PREFIX}/{cert.qr_id}.pdf"
            pdf_s3_link = CertificateService._upload_to_s3(pdf_path, pdf_key)

        # Persist the DOCX key as the stable s3_link (PDF can be re-derived from qr_id)
        cert.s3_link = docx_key
        db.session.commit()

        # Email: one message with DOCX always attached; PDF too if available
        today = date.today()
        try:
            valid_until = today.replace(year=today.year + 5).strftime("%B %d, %Y")
        except ValueError:
            valid_until = date(today.year + 5, 2, 28).strftime("%B %d, %Y")

        email_subject = f"Certificate of Appreciation \u2014 {course_code} ({academic_year})"
        email_body = f"""
<p>Dear {author_name},</p>

<p>On behalf of the institution, we are pleased to present you with a
<strong>Certificate of Appreciation</strong> in recognition of your valuable contribution
in developing an instructional material. The details of your certificate are as follows:</p>

<table style="border-collapse:collapse;font-size:14px;margin:12px 0;">
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">Certificate ID</td>
    <td style="padding:5px 0;">{cert.qr_id}</td>
  </tr>
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">Subject Code</td>
    <td style="padding:5px 0;">{course_code}</td>
  </tr>
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">Subject Title</td>
    <td style="padding:5px 0;">{course_title}</td>
  </tr>
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">Program</td>
    <td style="padding:5px 0;">{program_name}</td>
  </tr>
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">College</td>
    <td style="padding:5px 0;">{college_name}</td>
  </tr>
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">Semester</td>
    <td style="padding:5px 0;">{semester}</td>
  </tr>
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">Academic Year</td>
    <td style="padding:5px 0;">{academic_year}</td>
  </tr>
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">Date Issued</td>
    <td style="padding:5px 0;">{date_issued}</td>
  </tr>
  <tr>
    <td style="padding:5px 20px 5px 0;font-weight:bold;white-space:nowrap;color:#555;">Valid Until</td>
    <td style="padding:5px 0;">{valid_until}</td>
  </tr>
</table>

<p>Your certificate is attached to this email in DOCX and PDF formats.
Please retain a copy for your personal records. A QR code is embedded in the
certificate for quick verification.</p>

<p>We truly appreciate your dedication and efforts in enhancing the quality of
instruction. Your work reflects a deep commitment to academic excellence and to
the growth of your students.</p>

<p>Congratulations, and thank you.</p>

<p>Sincerely,<br>
<strong>Instructional Materials Management System (IMMS)</strong></p>
"""
        attachments = [(open(temp_output.name, 'rb').read(), f"{cert.qr_id}.docx")]
        if pdf_path and os.path.exists(pdf_path):
            attachments.append((open(pdf_path, 'rb').read(), f"{cert.qr_id}.pdf"))
        EmailService.send_files_to_recipients(
            user.email,
            attachments,
            subject=email_subject,
            html_body=email_body,
        )

        # Cleanup temp files
        os.remove(temp_output.name)
        if pdf_path:
            try:
                shutil.rmtree(os.path.dirname(pdf_path), ignore_errors=True)
            except Exception:
                pass

        return {
            'qr_id': cert.qr_id,
            'user_id': user.id,
            'author_name': author_name,
            's3_link': pdf_s3_link,           # PDF presigned URL, or None if conversion failed
            's3_link_docx': docx_s3_link,     # DOCX presigned URL, always present
        }

    @staticmethod
    def _build_author_name(user):
        """Build full display name for certificate text."""
        parts = [user.first_name]
        if user.middle_name:
            parts.append(user.middle_name)
        parts.append(user.last_name)
        return " ".join(p for p in parts if p)

    @staticmethod
    def _build_replacement_maps(
        college_name,
        course_code,
        course_title,
        author_rank,
        author_name,
        author_rank_and_name,
        program_name,
        semester_label,
        academic_year,
        date_issued,
        course_code_and_title,
        validity_duration,
    ):
        """Return literal and regex replacement maps for template substitution."""
        literal_replacements = {
            '{{COLLEGE_NAME}}': college_name,
            '{{COURSE_CODE}}': course_code,
            '{{COURSE_TITLE}}': course_title,
            '{{AUTHOR_RANK}}': author_rank,
            '{{AUTHOR_NAME}}': author_name,
            '{{AUTHOR_RANK_AND_NAME}}': author_rank_and_name,
            '{{PROGRAM_NAME}}': program_name,
            '{{SEMESTER}}': semester_label,
            '{{ACADEMIC_YEAR}}': academic_year,
            '{{DATE_ISSUED}}': date_issued,
            '{{COURSE_CODE_AND_TITLE}}': course_code_and_title,
            '{{IM_VALIDITY_DURATION}}': validity_duration,

            # Legacy placeholders used in earlier certificate templates.
            'Name of the College (NOC)': college_name,
            'Course Code: Course Title': course_code_and_title,
            'Rank and Name of Professor': author_rank_and_name,
            "Name of the Bachelor's Program (NOP)": program_name,
            "Name of the Bachelor’s Program (NOP)": program_name,
        }

        regex_replacements = [
            (r"Name\s+of\s+the\s+College\s*\(\s*NOC\s*\)", college_name),
            (r"Course\s+Code\s*:\s*Course\s+Title", course_code_and_title),
            (r"Rank\s+and\s+Name\s+of\s+Professor", author_rank_and_name),
            (r"Name\s+of\s+the\s+Bachelor(?:'|’)?s\s+Program\s*\(\s*NOP\s*\)", program_name),
        ]

        return literal_replacements, regex_replacements

    @staticmethod
    def _apply_template_replacements(
        doc,
        replacements,
        regex_replacements,
        semester_label,
        academic_year,
        date_issued,
        validity_duration,
    ):
        """Apply all replacement strategies across the whole document."""
        for paragraph in CertificateService._iter_all_paragraphs(doc):
            CertificateService._replace_text(paragraph, replacements)
            for pattern, replacement in regex_replacements:
                CertificateService._replace_regex_across_runs(
                    paragraph,
                    pattern,
                    replacement,
                    flags=re.IGNORECASE,
                )
            CertificateService._replace_legacy_patterns(
                paragraph,
                semester_label=semester_label,
                academic_year=academic_year,
                date_issued=date_issued,
                validity_duration=validity_duration,
            )
    
    @staticmethod
    def _replace_text(paragraph, replacements):
        """Replace known placeholders in a paragraph while preserving run style."""
        for key, value in replacements.items():
            CertificateService._replace_literal_across_runs(
                paragraph,
                str(key),
                str(value),
            )

    @staticmethod
    def _replace_legacy_patterns(paragraph, semester_label, academic_year, date_issued, validity_duration):
        """Handle older template phrasings that do not use {{TOKEN}} placeholders."""
        # Example: "1st semester" -> "2nd Semester"
        CertificateService._replace_regex_across_runs(
            paragraph,
            r"\b[1-4]\s*(?:st|nd|rd|th)\s+semester\b",
            semester_label,
            flags=re.IGNORECASE,
        )

        # Example: "Academic Year 2024 – 2025" -> "Academic Year 2026-2027"
        CertificateService._replace_regex_across_runs(
            paragraph,
            r"Academic\s+Year\s+\d{4}\s*[\-\u2013]\s*\d{4}",
            f"Academic Year {academic_year}",
            flags=re.IGNORECASE,
        )

        # Example date in narrative: "October 15, 2025"
        if 'issued on' in paragraph.text.lower():
            CertificateService._replace_regex_across_runs(
                paragraph,
                r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}\b",
                date_issued,
                flags=re.IGNORECASE,
            )

        # Example: "one academic year" -> dynamic validity duration.
        CertificateService._replace_regex_across_runs(
            paragraph,
            r"\bone\s+academic\s+year\b",
            validity_duration,
            flags=re.IGNORECASE,
        )

    @staticmethod
    def _iter_paragraphs_in_table(table):
        """Yield all paragraphs in a table, including nested tables."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested_table in cell.tables:
                    yield from CertificateService._iter_paragraphs_in_table(nested_table)

    @staticmethod
    def _iter_all_paragraphs(doc):
        """Yield paragraphs from body, tables, and section headers/footers."""
        for paragraph in doc.paragraphs:
            yield paragraph

        for table in doc.tables:
            yield from CertificateService._iter_paragraphs_in_table(table)

        for section in doc.sections:
            containers = [section.header, section.footer]
            for container in containers:
                for paragraph in container.paragraphs:
                    yield paragraph
                for table in container.tables:
                    yield from CertificateService._iter_paragraphs_in_table(table)

    @staticmethod
    def _replace_literal_across_runs(paragraph, old_text, new_text):
        """Replace a literal substring even when it is split across multiple runs."""
        if not old_text:
            return False

        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text or old_text not in full_text:
            return False

        matches = list(re.finditer(re.escape(old_text), full_text))
        if not matches:
            return False

        return CertificateService._replace_matches_across_runs(
            paragraph,
            matches,
            lambda _m: new_text,
        )

    @staticmethod
    def _replace_regex_across_runs(paragraph, pattern, replacement, flags=0):
        """Regex replace across runs while preserving the surrounding run formatting."""
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            return False

        matches = list(re.finditer(pattern, full_text, flags=flags))
        if not matches:
            return False

        if callable(replacement):
            replacement_fn = replacement
        else:
            replacement_fn = lambda _m: str(replacement)

        return CertificateService._replace_matches_across_runs(
            paragraph,
            matches,
            replacement_fn,
        )

    @staticmethod
    def _replace_matches_across_runs(paragraph, matches, replacement_fn):
        """Apply matched-span replacements by editing run text ranges in reverse order."""
        if not paragraph.runs:
            return False

        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            return False

        # Capture run boundaries once and replace from right-to-left so offsets remain valid.
        boundaries = []
        cursor = 0
        for idx, run in enumerate(paragraph.runs):
            run_len = len(run.text)
            boundaries.append((cursor, cursor + run_len, idx))
            cursor += run_len

        def locate(char_pos):
            for start, end, run_idx in boundaries:
                if start <= char_pos < end:
                    return run_idx, char_pos - start
            return None, None

        changed = False
        for match in reversed(matches):
            start, end = match.span()
            if start >= end:
                continue

            start_run_idx, start_off = locate(start)
            end_run_idx, end_off = locate(end - 1)
            if start_run_idx is None or end_run_idx is None:
                continue

            replacement_text = str(replacement_fn(match))

            if start_run_idx == end_run_idx:
                run = paragraph.runs[start_run_idx]
                run_text = run.text
                run.text = run_text[:start_off] + replacement_text + run_text[end_off + 1:]
            else:
                start_run = paragraph.runs[start_run_idx]
                end_run = paragraph.runs[end_run_idx]
                prefix = start_run.text[:start_off]
                suffix = end_run.text[end_off + 1:]
                start_run.text = prefix + replacement_text + suffix
                for i in range(start_run_idx + 1, end_run_idx + 1):
                    paragraph.runs[i].text = ""

            changed = True

        return changed

    @staticmethod
    def _format_semester_label(semester):
        """Normalize semester value for certificate wording."""
        if not semester:
            return "N/A"

        value = str(semester).strip()
        lowered = value.lower()

        if ('1' in lowered and 'sem' in lowered) or 'first sem' in lowered:
            return "1st Semester"
        if ('2' in lowered and 'sem' in lowered) or 'second sem' in lowered:
            return "2nd Semester"
        if ('3' in lowered and 'sem' in lowered) or 'third sem' in lowered:
            return "3rd Semester"

        if 'semester' in lowered:
            return value.replace('semester', 'Semester')

        return value

    @staticmethod
    def _format_validity_duration(validity):
        """Map IM validity data into a sentence-friendly duration phrase."""
        default_duration = "one academic year"
        if validity is None:
            return default_duration

        text = str(validity).strip()
        if not text:
            return default_duration

        lowered = text.lower()
        if 'academic year' in lowered or 'semester' in lowered or 'year' in lowered:
            return text

        # Common form in this codebase: a 4-digit year (e.g., 2026) means one AY.
        if re.fullmatch(r"\d{4}", text):
            return default_duration

        # If explicitly numeric and short, treat as count of academic years.
        if re.fullmatch(r"\d{1,2}", text):
            count = int(text)
            if count <= 1:
                return default_duration
            return f"{count} academic years"

        # If given a year range, derive a duration count.
        range_match = re.fullmatch(r"(\d{4})\s*[\-\u2013]\s*(\d{4})", text)
        if range_match:
            start_year = int(range_match.group(1))
            end_year = int(range_match.group(2))
            span = max(1, end_year - start_year)
            if span <= 1:
                return default_duration
            return f"{span} academic years"

        return text
    
    @staticmethod
    def _generate_qr_code(data):
        """Generate QR code image"""
        qr = qrcode.QRCode(version=1, box_size=10, border=2)
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Convert to bytes
        img_bytes = BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        return img_bytes
    
    @staticmethod
    def _add_qr_to_document(doc, qr_img):
        """Add QR code by placeholder; fallback to insertion before approval/signature area."""

        # Prefer explicit placeholder where available.
        for paragraph in CertificateService._iter_all_paragraphs(doc):
            if CertificateService.QR_PLACEHOLDER in paragraph.text:
                paragraph.clear()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = paragraph.add_run()
                run.add_picture(qr_img, width=Inches(CertificateService.QR_INLINE_WIDTH_INCHES))
                return

        # Fallback: place QR as floating image so it won't reflow text into another page.
        fallback_paragraph = CertificateService._find_qr_fallback_paragraph(doc)
        if fallback_paragraph is not None:
            if CertificateService._add_floating_qr_at_paragraph(
                fallback_paragraph,
                qr_img,
                width_inches=CertificateService.QR_FLOATING_WIDTH_INCHES,
            ):
                return

        # Final fallback: append small inline QR if floating mode fails.
        qr_paragraph = doc.add_paragraph()
        qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        qr_paragraph.add_run().add_picture(qr_img, width=Inches(1.1))

    @staticmethod
    def _find_qr_fallback_paragraph(doc):
        """Choose a stable paragraph near the approval section for floating QR anchoring."""
        approved_idx = None
        for idx, paragraph in enumerate(doc.paragraphs):
            if 'Approved by' in paragraph.text:
                approved_idx = idx
                break

        if approved_idx is None:
            return doc.paragraphs[-1] if doc.paragraphs else None

        # Prefer an empty paragraph above approval to avoid touching text layout.
        for idx in range(approved_idx - 1, -1, -1):
            if not doc.paragraphs[idx].text.strip():
                return doc.paragraphs[idx]

        return doc.paragraphs[approved_idx]

    @staticmethod
    def _add_floating_qr_at_paragraph(paragraph, qr_img, width_inches):
        """Add QR as floating drawing anchored to a paragraph to avoid content reflow."""
        try:
            run = paragraph.add_run()
            run.add_picture(qr_img, width=Inches(width_inches))
            drawing_nodes = run._r.xpath('./w:drawing')
            if not drawing_nodes:
                return False
            drawing = drawing_nodes[0]
            inline_nodes = drawing.xpath('./wp:inline')
            if not inline_nodes:
                return False

            inline = inline_nodes[0]
            anchor = CertificateService._inline_to_anchor_xml(inline)
            drawing.remove(inline)
            drawing.append(anchor)
            return True
        except Exception:
            return False

    @staticmethod
    def _inline_to_anchor_xml(inline):
        """Convert a wp:inline node into wp:anchor for non-flow floating placement."""
        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '0')
        anchor.set('distR', '0')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251659264')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')

        simple_pos = OxmlElement('wp:simplePos')
        simple_pos.set('x', '0')
        simple_pos.set('y', '0')
        anchor.append(simple_pos)

        position_h = OxmlElement('wp:positionH')
        position_h.set('relativeFrom', 'margin')
        align_h = OxmlElement('wp:align')
        align_h.text = 'right'
        position_h.append(align_h)
        anchor.append(position_h)

        position_v = OxmlElement('wp:positionV')
        position_v.set('relativeFrom', 'paragraph')
        pos_offset_v = OxmlElement('wp:posOffset')
        pos_offset_v.text = '0'
        position_v.append(pos_offset_v)
        anchor.append(position_v)

        extent = inline.find(qn('wp:extent'))
        if extent is not None:
            anchor.append(extent)

        effect_extent = OxmlElement('wp:effectExtent')
        effect_extent.set('l', '0')
        effect_extent.set('t', '0')
        effect_extent.set('r', '0')
        effect_extent.set('b', '0')
        anchor.append(effect_extent)

        wrap_none = OxmlElement('wp:wrapNone')
        anchor.append(wrap_none)

        doc_pr = inline.find(qn('wp:docPr'))
        if doc_pr is not None:
            anchor.append(doc_pr)

        cnv = inline.find(qn('wp:cNvGraphicFramePr'))
        if cnv is not None:
            anchor.append(cnv)

        graphic = inline.find(qn('a:graphic'))
        if graphic is not None:
            anchor.append(graphic)

        return anchor
    
    @staticmethod
    def _upload_to_s3(file_path, s3_key):
        """Upload a file to S3 and return a presigned URL valid for 7 days."""
        bucket_name = os.getenv('AWS_BUCKET_NAME')
        s3 = boto3.client('s3')
        s3.upload_file(file_path, bucket_name, s3_key)
        presigned_url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': bucket_name, 'Key': s3_key},
            ExpiresIn=604800  # 7 days
        )
        return presigned_url

    @staticmethod
    def _convert_docx_to_pdf(docx_path):
        """Convert a DOCX file to PDF.
        Tries docx2pdf first (uses Word on Windows, LibreOffice on Linux/macOS),
        then falls back to a raw LibreOffice subprocess.
        Returns the path to the generated PDF, or None if both methods fail.
        """
        out_dir = tempfile.mkdtemp()
        base = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(out_dir, base + '.pdf')

        # --- Method 1: docx2pdf (cross-platform) ---
        try:
            import pythoncom  # type: ignore
            from docx2pdf import convert as d2p_convert  # type: ignore
            pythoncom.CoInitialize()
            try:
                d2p_convert(docx_path, pdf_path)
            finally:
                pythoncom.CoUninitialize()
            if os.path.exists(pdf_path):
                return pdf_path
        except Exception as e:
            print(f"[cert] docx2pdf failed ({e}), trying LibreOffice…")

        # --- Method 2: LibreOffice headless subprocess ---
        try:
            subprocess.run(
                [
                    'libreoffice', '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', out_dir,
                    docx_path,
                ],
                check=True,
                capture_output=True,
                timeout=120,
            )
            lo_pdf = os.path.join(out_dir, base + '.pdf')
            if os.path.exists(lo_pdf):
                return lo_pdf
        except Exception as e:
            print(f"[cert] LibreOffice fallback also failed ({e})")

        shutil.rmtree(out_dir, ignore_errors=True)
        return None  # caller must handle gracefully

    @staticmethod
    def _key_exists_in_s3(key):
        """Return True if the given S3 key exists in the bucket."""
        bucket_name = os.getenv('AWS_BUCKET_NAME')
        try:
            boto3.client('s3').head_object(Bucket=bucket_name, Key=key)
            return True
        except Exception:
            return False

    @staticmethod
    def _try_presign_pdf(qr_id):
        """Return a presigned URL for the PDF of this cert, or None if it doesn't exist yet."""
        key = f"{CertificateService.GENERATED_CERTIFICATES_PREFIX}/{qr_id}.pdf"
        if not CertificateService._key_exists_in_s3(key):
            return None
        bucket_name = os.getenv('AWS_BUCKET_NAME')
        return boto3.client('s3').generate_presigned_url(
            'get_object',
            Params={'Bucket': bucket_name, 'Key': key},
            ExpiresIn=604800
        )

    @staticmethod
    def _resolve_s3_link(raw_link):
        """Ensure raw_link is a usable HTTPS URL.
        - Already a URL → return as-is.
        - Legacy DOCX key → rewrite to .pdf key before presigning.
        - Other S3 key → presign directly.
        """
        if not raw_link:
            return raw_link
        if raw_link.startswith('http://') or raw_link.startswith('https://'):
            return raw_link
        bucket_name = os.getenv('AWS_BUCKET_NAME')
        key = raw_link
        if key.endswith('.docx'):
            key = key[:-5] + '.pdf'
        try:
            s3 = boto3.client('s3')
            return s3.generate_presigned_url(
                'get_object',
                Params={'Bucket': bucket_name, 'Key': key},
                ExpiresIn=604800
            )
        except Exception:
            return raw_link

    @staticmethod
    def _resolve_docx_link(qr_id):
        """Generate a fresh presigned URL for the DOCX of a given cert (by qr_id)."""
        bucket_name = os.getenv('AWS_BUCKET_NAME')
        key = f"{CertificateService.GENERATED_CERTIFICATES_PREFIX}/{qr_id}.docx"
        try:
            s3 = boto3.client('s3')
            return s3.generate_presigned_url(
                'get_object',
                Params={'Bucket': bucket_name, 'Key': key},
                ExpiresIn=604800
            )
        except Exception:
            return None
